from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\dev\SiCoSe")
OUTPUT_DIR = ROOT / "output" / "word"
ASSET_DIR = OUTPUT_DIR / "consultorio_web_assets"
DOCX_PATH = OUTPUT_DIR / "AUDIT-1-Consultorio-Digital-SaaS-evidencia-real.docx"

SKILL_SCRIPTS = Path(
    r"C:\Users\SAM\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BASE_FONT = "Calibri"
ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(85, 85, 85)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"

IMAGE_FILES = [
    "home-top.png",
    "home-services.png",
    "login-page.png",
    "login-verifying.png",
    "patient-page.png",
]


def set_run_font(
    run,
    *,
    name: str = BASE_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell(cell) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10,
    color: RGBColor = INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def require_assets() -> dict[str, Path]:
    missing = [name for name in IMAGE_FILES if not (ASSET_DIR / name).exists()]
    if missing:
        raise FileNotFoundError(f"Missing evidence images: {', '.join(missing)}")
    return {name: ASSET_DIR / name for name in IMAGE_FILES}


def style_document(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_DARK, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    left = paragraph.add_run("AUDIT-1 | Consultorio Digital SaaS")
    set_run_font(left, size=9, color=RGBColor(99, 110, 125))
    paragraph.add_run("\t")
    right = paragraph.add_run("Página ")
    set_run_font(right, size=9, color=RGBColor(99, 110, 125))
    add_page_number_field(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size=9, color=RGBColor(99, 110, 125))


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    set_paragraph_format(kicker, before=0, after=2, line_spacing=1.0)
    run = kicker.add_run("AUDIT-1 | EVALUACIÓN MVP - SITIO EN VIVO")
    set_run_font(run, size=10, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    set_paragraph_format(title, before=0, after=4, line_spacing=1.0)
    run = title.add_run("Consultorio Digital SaaS")
    set_run_font(run, size=26, color=ACCENT_DARK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, before=0, after=10, line_spacing=1.05)
    run = subtitle.add_run(
        "Evaluación funcional con evidencia real del dominio consultorio-digital-saas.vercel.app"
    )
    set_run_font(run, size=13, color=MUTED)


def add_label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.7, 4.8], 9360)
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.rows[index].cells[0], label, bold=True, size=10, color=ACCENT_DARK)
        shade_cell(table.rows[index].cells[0], HEADER_FILL)
        set_cell_text(table.rows[index].cells[1], value, size=10, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_light_callout(doc: Document, headline: str, body: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=8, line_spacing=1.05)
    shade_paragraph(paragraph, LIGHT_FILL)
    paragraph.paragraph_format.left_indent = Inches(0.05)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    run = paragraph.add_run(f"{headline} ")
    set_run_font(run, size=10.75, color=INK, bold=True)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.75, color=INK)


def add_three_column_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, str, str]],
    widths: list[float],
    *,
    header_size: float = 10,
    body_size: float = 9.5,
) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=header_size, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[index], text, size=body_size, color=INK, align=alignment)
    apply_table_geometry(table, column_widths_from_weights(widths, 9360), table_width_dxa=9360, indent_dxa=120)


def add_image_figure(doc: Document, image_path: Path, caption: str) -> None:
    figure = doc.add_paragraph()
    set_paragraph_format(figure, before=0, after=2, line_spacing=1.0)
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.add_run().add_picture(str(image_path), width=Inches(6.3))

    caption_paragraph = doc.add_paragraph()
    set_paragraph_format(caption_paragraph, before=0, after=10, line_spacing=1.0)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9.25, color=MUTED, italic=True)


def add_cover_page(doc: Document) -> None:
    add_title_block(doc)

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=10, line_spacing=1.08)
    run = intro.add_run(
        "El sitio responde en vivo y su navegación principal es clara. "
        "Esta entrega documenta evidencia real del dominio, con pruebas de acceso al portal del personal y al portal del paciente."
    )
    set_run_font(run, size=11, color=INK)

    add_label_value_table(
        doc,
        [
            ("Link evaluado", "https://consultorio-digital-saas.vercel.app/"),
            ("Repositorio", "https://github.com/SoyAri/consultorio-digital"),
            (
                "Equipo evaluador",
                "David Aguilar Rodriguez | Benkis Carbajal Henandez | Cesar Gaspar Pacheco | Samuel Jonathan Trujillo Bolaños",
            ),
            ("Fecha", "18 de junio de 2026"),
            ("Entorno de prueba", "Browser integrado Codex, pantalla 1280x720, Windows"),
            ("Estado", "Validacion viva parcial con evidencia real"),
        ],
    )

    add_light_callout(
        doc,
        "Alcance:",
        "se verificaron carga de portada, navegación a secciones, apertura del login del personal, estado de verificación del acceso y portal del paciente. "
        "No se enviaron SMS reales ni se usaron datos sensibles.",
    )


def add_summary_page(doc: Document) -> None:
    heading = doc.add_heading("1. Resumen ejecutivo", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "El producto se presenta como una clínica dental con dos entradas claras: pacientes y personal. "
        "La primera impresión es buena: la portada carga, el mensaje se entiende en segundos y los llamados a la acción son visibles."
    )
    set_run_font(run, size=11, color=INK)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "En pruebas vivas, la navegación funciona en la capa pública y las rutas de acceso existen. "
        "La falla relevante aparece al intentar autenticar al personal: después de enviar credenciales de ensayo, la interfaz queda atascada en \"Verificando...\" y no completa el flujo."
    )
    set_run_font(run, size=11, color=INK)

    add_three_column_table(
        doc,
        ["Criterio", "Resultado", "Lectura QA"],
        [
            ("Funciona", "Sí", "La portada, navegación y formularios visibles responden."),
            ("Flujo completo", "No", "El login del personal no finaliza en la ventana observada."),
            ("Se entiende", "Sí", "El producto comunica bien qué hace y a quién sirve."),
            ("Bugs críticos", "Sí", "El acceso del personal bloquea la ruta principal."),
            ("Score binario", "3/4", "Valoración positiva parcial, no liberable."),
            ("Dictamen", "Parcial", "Se ve real, pero aún no cierra el flujo principal."),
        ],
        [1.5, 1.0, 4.0],
    )


def add_methodology_page(doc: Document) -> None:
    heading = doc.add_heading("2. Metodología de prueba", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=8, line_spacing=1.08)
    run = intro.add_run(
        "La validación se hizo navegando el sitio en vivo desde el navegador integrado. "
        "Se usaron datos de ensayo no reales para observar estados de UI, sin transmitir información sensible ni completar acciones que implicaran costo externo."
    )
    set_run_font(run, size=11, color=INK)

    add_three_column_table(
        doc,
        ["Flujo", "Resultado", "Evidencia"],
        [
            ("Portada", "Carga limpia y CTA visibles", "home-top.png"),
            ("Servicios", "El ancla lleva a la sección correcta", "home-services.png"),
            ("Personal", "La ruta /login/equipo abre formulario", "login-page.png"),
            ("Login staff", "El estado cambia a Verificando...", "login-verifying.png"),
            ("Paciente", "La ruta /login/paciente abre portal", "patient-page.png"),
        ],
        [1.5, 3.8, 1.2],
        body_size=9.4,
    )

    add_light_callout(
        doc,
        "Nota de seguridad:",
        "Se evitó presionar acciones que pudieran enviar SMS, abrir costos externos o comprometer datos reales. "
        "La auditoría priorizó evidencia segura y reproducible.",
    )


def add_findings_page(doc: Document) -> None:
    heading = doc.add_heading("3. Hallazgos y bugs observados", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "Los hallazgos no apuntan a una pantalla simulada. Hay rutas reales, cambios de estado y formularios funcionales. "
        "El problema está en el cierre del flujo: el login del personal no termina y deja al usuario sin retroalimentación útil."
    )
    set_run_font(run, size=11, color=INK)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Hallazgo", "Severidad", "Impacto QA"]
    widths = column_widths_from_weights([2.0, 1.0, 3.5], 9360)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)

    bugs = [
        (
            "Login del personal bloqueado",
            "Alta",
            "Tras escribir qa@demo.com / Demo1234 y pulsar Entrar al sistema, la pantalla queda en Verificando... y no completa el acceso.",
        ),
        (
            "Sin feedback visible de error",
            "Media",
            "No aparece mensaje de fallo, timeout ni reintento; el usuario queda sin señal clara de lo que ocurrió.",
        ),
        (
            "Flujo de SMS no validado",
            "Media",
            "El portal del paciente habilita el botón con un teléfono de 10 dígitos, pero completar el envío requeriría un SMS real.",
        ),
    ]
    for bug in bugs:
        row = table.add_row()
        set_cell_text(row.cells[0], bug[0], size=9.4, color=INK)
        set_cell_text(row.cells[1], bug[1], size=9.4, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], bug[2], size=9.2, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)

    add_light_callout(
        doc,
        "Lectura de riesgo:",
        "el sitio sí opera como una aplicación real en su capa pública, pero la autenticación del personal sigue siendo el punto que impide cerrar el MVP.",
    )


def add_final_page(doc: Document) -> None:
    heading = doc.add_heading("4. Evaluación final", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "Resultado resumido: la experiencia pública está bien planteada y el producto se entiende sin fricción. "
        "Sin embargo, el bloqueo del login del personal afecta el corazón operativo del MVP."
    )
    set_run_font(run, size=11, color=INK)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.8, 4.7], 9360)
    headers = ["Campo", "Valor"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)

    rows = [
        ("Funciona", "Sí"),
        ("Flujo completo", "No"),
        ("Se entiende", "Sí"),
        ("Bugs críticos", "Sí"),
        ("Score binario", "3/4"),
        ("Dictamen", "Funcional parcial; no liberable hasta resolver login"),
    ]
    for label, value in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, size=9.8, color=INK, bold=True)
        set_cell_text(row.cells[1], value, size=9.8, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)

    add_light_callout(
        doc,
        "Cierre QA:",
        "si se corrige el cierre del login del personal y se confirma el manejo de errores, este reporte puede pasar de parcial a funcional en la siguiente revisión.",
    )


def add_evidence_intro(doc: Document) -> None:
    heading = doc.add_heading("5. Evidencia real del sitio", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=10, line_spacing=1.08)
    run = paragraph.add_run(
        "Las siguientes capturas provienen del dominio en vivo y muestran el recorrido realizado durante la auditoría."
    )
    set_run_font(run, size=11, color=INK)


def add_evidence_page(doc: Document, image_path: Path, heading_text: str, caption: str) -> None:
    doc.add_page_break()
    heading = doc.add_heading(heading_text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_image_figure(doc, image_path, caption)


def build_doc() -> Path:
    assets = require_assets()

    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0])
    doc.settings.odd_and_even_pages_header_footer = False
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    add_cover_page(doc)
    doc.add_page_break()
    add_footer(doc.sections[0])

    add_summary_page(doc)
    doc.add_page_break()
    add_methodology_page(doc)
    doc.add_page_break()
    add_findings_page(doc)
    doc.add_page_break()
    add_final_page(doc)
    doc.add_page_break()
    add_evidence_intro(doc)

    evidence_specs = [
        ("home-top.png", "Evidencia 1 - Portada del sitio", "A1 - Captura del hero principal y los accesos públicos visibles en la portada."),
        ("home-services.png", "Evidencia 2 - Sección de servicios", "A2 - Vista de la sección de servicios tras usar el ancla de navegación."),
        ("login-page.png", "Evidencia 3 - Login del personal", "A3 - Ruta /login/equipo con formulario activo y botón habilitado."),
        ("login-verifying.png", "Evidencia 4 - Login en verificación", "A4 - El intento de acceso queda en Verificando... y no completa el flujo."),
        ("patient-page.png", "Evidencia 5 - Portal del paciente", "A5 - Ruta /login/paciente con el campo de teléfono validado y el botón activo."),
    ]

    for image_name, heading_text, caption in evidence_specs:
        add_evidence_page(doc, assets[image_name], heading_text, caption)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
