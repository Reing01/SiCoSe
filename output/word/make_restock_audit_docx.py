from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS = Path(
    r"C:\Users\SAM\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ROOT = Path(r"C:\dev\SiCoSe")
ASSET_DIR = ROOT / "output" / "pdf" / "assets"
OUTPUT_DIR = ROOT / "output" / "word"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = (
    OUTPUT_DIR
    / "AUDIT-1-David-Aguilar-Rodriguez-Benkis-Carbajal-Henandez-Cesar-Gaspar-Pacheco-Samuel-Jonathan-Trujillo-Bolanos-to-Arturocs160-ReStock-SaaS.docx"
)


BASE_FONT = "Calibri"
ACCENT_BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "F2F4F7"
SOFT_BLUE = "E8EEF5"
WHITE = "FFFFFF"


def set_run_font(
    run,
    *,
    name: str = BASE_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph_format(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell(cell) -> None:
    cell.text = ""
    if cell.paragraphs:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10,
    color: RGBColor = INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text or url
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def style_document(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = BASE_FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)

    heading1 = styles["Heading 1"]
    heading1.font.name = BASE_FONT
    heading1._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = DARK_BLUE
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    heading2 = styles["Heading 2"]
    heading2.font.name = BASE_FONT
    heading2._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    heading2._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = ACCENT_BLUE
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.keep_with_next = True

    heading3 = styles["Heading 3"]
    heading3.font.name = BASE_FONT
    heading3._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    heading3._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = DARK_BLUE
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )

    run = paragraph.add_run("AUDIT-1 | ReStock SaaS")
    set_run_font(run, size=9, color=RGBColor(99, 110, 125))
    run.add_tab()

    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=RGBColor(99, 110, 125))
    add_page_number_field(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size=9, color=RGBColor(99, 110, 125))


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    set_paragraph_format(kicker, before=0, after=2, line_spacing=1.0)
    run = kicker.add_run("AUDIT-1 | EVALUACION DE MVP (SEMANA 7)")
    set_run_font(run, size=10, color=ACCENT_BLUE, bold=True)

    title = doc.add_paragraph()
    set_paragraph_format(title, before=0, after=4, line_spacing=1.0, keep_with_next=True)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("ReStock SaaS")
    set_run_font(run, size=26, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, before=0, after=10, line_spacing=1.05)
    run = subtitle.add_run("Revision funcional del MVP del repositorio Arturocs160/ReStock-SaaS")
    set_run_font(run, size=13, color=MUTED)


def add_label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.6, 4.9], 9360)
    for row_index, (label, value) in enumerate(rows):
        set_cell_text(table.rows[row_index].cells[0], label, bold=True, size=10, color=DARK_BLUE)
        set_cell_shading(table.rows[row_index].cells[0], HEADER_FILL)
        value_cell = table.rows[row_index].cells[1]
        clear_cell(value_cell)
        paragraph = value_cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if isinstance(value, tuple):
            text, url = value
            add_hyperlink(paragraph, url, text)
            for run in paragraph.runs:
                set_run_font(run, size=10, color=INK)
        else:
            run = paragraph.add_run(value)
            set_run_font(run, size=10, color=INK)
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    return table


def add_binary_score_callout(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=10, line_spacing=1.05)
    set_paragraph_shading(paragraph, LIGHT_FILL)
    paragraph.paragraph_format.left_indent = Inches(0.05)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    run = paragraph.add_run("Resultado final: funcional. ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = paragraph.add_run(
        "El flujo principal responde con comportamiento real, la informacion del producto se entiende con facilidad y el inventario persiste. "
        "Se observaron modulos secundarios simulados, pero no se detectaron fallas criticas que bloqueen la validacion del MVP."
    )
    set_run_font(run, size=11, color=INK)


def add_score_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([2.1, 1.0, 3.4], 9360)

    headers = ["Criterio", "Resultado", "Evidencia"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=10, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, SOFT_BLUE)

    data = [
        ("Funciona", "Si", "E1-E6"),
        ("Flujo completo", "Si", "E1-E6"),
        ("Se entiende", "Si", "E1-E3"),
        ("Bugs criticos", "No", "E7-E10"),
    ]
    for row_data in data:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=9.5, color=INK, align=align)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_environment_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.6, 4.9], 9360)

    headers = ["Parametro", "Detalle"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(table.rows[0].cells[idx], SOFT_BLUE)

    rows = [
        ("Link evaluado", "https://www.restock.website/"),
        ("Repositorio", "Arturocs160/ReStock-SaaS"),
        ("Credenciales demo", "admin@demo.com / Demo123@"),
        ("Navegador de prueba", "Navegador integrado de Codex"),
        ("Fecha de revision", "18 de junio de 2026"),
    ]
    for label, detail in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True, size=9.5, color=DARK_BLUE)
        set_cell_shading(row.cells[0], HEADER_FILL)
        set_cell_text(row.cells[1], detail, size=9.5, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_criteria_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([2.1, 1.0, 3.4], 9360)

    headers = ["Criterio binario", "Resultado", "Evidencia"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], SOFT_BLUE)

    rows = [
        ("Funciona", "Si", "E1-E6"),
        ("Flujo completo", "Si", "E1-E6"),
        ("Se entiende", "Si", "E1-E3"),
        ("Bugs criticos", "No", "E7-E10"),
    ]
    for label, result, evidence in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, size=9.5, color=INK)
        set_cell_text(row.cells[1], result, size=9.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], evidence, size=9.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_functional_matrix(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([0.6, 2.0, 0.9, 1.0, 2.0], 9360)

    headers = ["ID", "Escenario", "Resultado", "Evidencia", "Observacion"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.5, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], SOFT_BLUE)

    rows = [
        ("T1", "Abrir landing page y validar carga inicial", "Aprobado", "E1", "El sitio carga y comunica claramente su propuesta."),
        ("T2", "Abrir login y revisar credenciales demo visibles", "Aprobado", "E2", "La interfaz muestra el acceso y la referencia de prueba."),
        ("T3", "Iniciar sesion con admin@demo.com / Demo123@", "Aprobado", "E3", "El acceso autentica y abre el dashboard operativo."),
        ("T4", "Revisar resumen operativo y navegacion principal", "Aprobado", "E3", "La experiencia central es coherente y navegable."),
        ("T5", "Abrir modal de alta de producto y guardar", "Aprobado", "E5", "El formulario aparece y permite registrar un producto."),
        ("T6", "Abrir modal de alta de lote y guardar", "Aprobado", "E6", "El lote modifica el stock y conserva el dato en la vista."),
        ("T7", "Recargar inventario y verificar persistencia", "Aprobado", "E4", "El inventario mantiene cambios luego del refresco."),
    ]
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in {0, 2, 3} else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=9, color=INK, align=align)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_bug_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.4, 3.2, 0.9, 1.0], 9360)

    headers = ["Ruta", "Hallazgo e impacto", "Severidad", "Evidencia"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.5, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], SOFT_BLUE)

    rows = [
        ("/dashboard/ventas", "La vista muestra contenido de demostracion y no expone un formulario operativo para registrar ventas.", "Media", "E7"),
        ("/dashboard/vencimientos", "La pantalla solo informa que es simulada; no entrega un listado accionable de proximos vencimientos.", "Media", "E8"),
        ("/dashboard/compras", "La ruta funciona como maqueta y no presenta flujo real de abastecimiento.", "Media", "E9"),
        ("/dashboard/configuracion", "No hay ajustes reales del sistema; la pagina solo replica una vista de demostracion.", "Baja", "E10"),
    ]
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in {2, 3} else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=9, color=INK, align=align)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_image_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    picture_paragraph = doc.add_paragraph()
    set_paragraph_format(picture_paragraph, before=0, after=2, line_spacing=1.0)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(5.85))

    caption_paragraph = doc.add_paragraph()
    set_paragraph_format(caption_paragraph, before=0, after=8, line_spacing=1.0)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=9.25, color=MUTED, italic=True)


def add_cover_page(doc: Document) -> None:
    add_title_block(doc)

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=10, line_spacing=1.08)
    run = intro.add_run("Objetivo: validar que el MVP carga, se puede usar, completa el flujo principal y presenta evidencias reales del comportamiento observado.")
    set_run_font(run, size=11, color=INK)

    rows = [
        ("Link evaluado", ("https://www.restock.website/", "https://www.restock.website/")),
        ("Repositorio", ("Arturocs160/ReStock-SaaS", "https://github.com/Arturocs160/ReStock-SaaS")),
        (
            "Equipo evaluador",
            "David Aguilar Rodriguez; Benkis Carbajal Henandez; Cesar Gaspar Pacheco; Samuel Jonathan Trujillo Bolaños",
        ),
        ("Fecha de revision", "18 de junio de 2026"),
    ]
    add_label_value_table(doc, rows)

    note = doc.add_paragraph()
    set_paragraph_format(note, before=6, after=0, line_spacing=1.05)
    set_paragraph_shading(note, LIGHT_FILL)
    note.paragraph_format.left_indent = Inches(0.05)
    note.paragraph_format.right_indent = Inches(0.05)
    run = note.add_run("Alcance: revision de carga, acceso, inventario real, persistencia de cambios y rutas secundarias simuladas.")
    set_run_font(run, size=10.5, color=INK)


def add_exec_summary_page(doc: Document) -> None:
    heading = doc.add_heading("Resumen ejecutivo", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_binary_score_callout(doc)
    add_criteria_table(doc)

    text = doc.add_paragraph()
    set_paragraph_format(text, before=8, after=4, line_spacing=1.08)
    run = text.add_run("Lectura QA: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = text.add_run(
        "la pantalla de inicio comunica el producto con claridad, el acceso demo funciona y el inventario muestra comportamiento real. "
        "La unica brecha relevante es que algunas rutas secundarias siguen funcionando como simulacion."
    )
    set_run_font(run, size=11, color=INK)

    text = doc.add_paragraph()
    set_paragraph_format(text, before=0, after=0, line_spacing=1.08)
    run = text.add_run("Recomendacion: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = text.add_run(
        "convertir las rutas de ventas, vencimientos, compras y configuracion en pantallas operativas o etiquetarlas de manera explicita como demo."
    )
    set_run_font(run, size=11, color=INK)


def add_methodology_page(doc: Document) -> None:
    heading = doc.add_heading("1. Alcance y metodologia de prueba", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=8, line_spacing=1.08)
    run = intro.add_run(
        "La validacion se ejecuto como una revision funcional de tipo smoke test, enfocada en carga inicial, autenticacion, operacion del inventario y verificacion visual de rutas secundarias."
    )
    set_run_font(run, size=11, color=INK)

    add_environment_table(doc)

    note = doc.add_paragraph()
    set_paragraph_format(note, before=8, after=4, line_spacing=1.08)
    run = note.add_run("Flujo probado: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = note.add_run(
        "entrada al sitio, lectura de la propuesta de valor, apertura de login, autenticacion demo, exploracion del dashboard y pruebas de CRUD en inventario."
    )
    set_run_font(run, size=11, color=INK)


def add_functional_page(doc: Document) -> None:
    heading = doc.add_heading("2. Matriz de pruebas funcionales", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=8, line_spacing=1.05)
    run = intro.add_run("Cobertura: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = intro.add_run(
        "se priorizaron los recorridos que un usuario real tocaria primero para confirmar que el MVP no esta solamente decorado."
    )
    set_run_font(run, size=11, color=INK)

    add_functional_matrix(doc)

    note = doc.add_paragraph()
    set_paragraph_format(note, before=8, after=0, line_spacing=1.05)
    run = note.add_run("Validacion de persistencia: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = note.add_run(
        "la alta de producto y la alta de lote cambian el stock visible, y el estado se conserva despues de recargar el inventario."
    )
    set_run_font(run, size=11, color=INK)


def add_bugs_page(doc: Document) -> None:
    heading = doc.add_heading("3. Hallazgos y bugs", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=8, line_spacing=1.05)
    run = intro.add_run("Observacion general: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = intro.add_run(
        "no se observaron errores bloqueantes en el flujo principal, pero si varias rutas secundarias que funcionan como simulacion."
    )
    set_run_font(run, size=11, color=INK)

    add_bug_table(doc)

    verdict = doc.add_paragraph()
    set_paragraph_format(verdict, before=8, after=0, line_spacing=1.05)
    set_paragraph_shading(verdict, LIGHT_FILL)
    verdict.paragraph_format.left_indent = Inches(0.05)
    verdict.paragraph_format.right_indent = Inches(0.05)
    run = verdict.add_run("Dictamen binario final: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = verdict.add_run("Funciona: Si. Flujo completo: Si. Se entiende: Si. Bugs criticos: No. Score: 4/4 = funcional.")
    set_run_font(run, size=11, color=INK)

    closing = doc.add_paragraph()
    set_paragraph_format(closing, before=4, after=0, line_spacing=1.05)
    run = closing.add_run("Siguiente mejora sugerida: ")
    set_run_font(run, size=11, color=INK, bold=True)
    run = closing.add_run(
        "desarrollar las rutas demo para convertirlas en pantallas operativas o rotularlas claramente como contenido de demostracion."
    )
    set_run_font(run, size=11, color=INK)


def add_appendix_page(doc: Document, title: str, image_pairs: list[tuple[str, str]]) -> None:
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=10, line_spacing=1.05)
    run = intro.add_run("Capturas reales tomadas durante la revision para respaldar la evaluacion.")
    set_run_font(run, size=11, color=INK)

    for caption, image_name in image_pairs:
        add_image_figure(doc, ASSET_DIR / image_name, caption)


def build_doc() -> Path:
    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0])
    doc.settings.odd_and_even_pages_header_footer = False
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    # First page remains clean; footer is added only for the main section.
    add_cover_page(doc)
    doc.add_page_break()
    add_footer(doc.sections[0])

    add_exec_summary_page(doc)
    doc.add_page_break()
    add_methodology_page(doc)
    doc.add_page_break()
    add_functional_page(doc)
    doc.add_page_break()
    add_bugs_page(doc)

    doc.add_page_break()
    add_appendix_page(
        doc,
        "Anexo fotografico 1 - Inicio y acceso",
        [
            ("E1 - Landing page real: la propuesta de valor carga correctamente y presenta la aplicacion.", "landing.png"),
            ("E2 - Login real: el formulario muestra las credenciales demo y la entrada al sistema.", "login-form.png"),
        ],
    )
    doc.add_page_break()
    add_appendix_page(
        doc,
        "Anexo fotografico 2 - Sesion y operacion",
        [
            ("E3 - Dashboard autenticado: la sesion demo abre el resumen operativo con navegacion real.", "dashboard.png"),
            ("E4 - Inventario real: los productos y lotes muestran stock persistente y alertas de bajo inventario.", "inventory-view.png"),
        ],
    )
    doc.add_page_break()
    add_appendix_page(
        doc,
        "Anexo fotografico 3 - CRUD de inventario",
        [
            ("E5 - Alta de producto: modal funcional para registrar un producto de prueba.", "product-modal.png"),
            ("E6 - Alta de lote: modal funcional para registrar lotes y actualizar stock.", "lot-modal.png"),
        ],
    )
    doc.add_page_break()
    add_appendix_page(
        doc,
        "Anexo fotografico 4 - Modulos simulados 1",
        [
            ("E7 - Ventas simuladas: la ruta no expone un formulario operativo.", "sales-simulada.png"),
            ("E8 - Vencimientos simulados: la ruta solo muestra contenido de demostracion.", "vencimientos-simulada.png"),
        ],
    )
    doc.add_page_break()
    add_appendix_page(
        doc,
        "Anexo fotografico 5 - Modulos simulados 2",
        [
            ("E9 - Compras simuladas: la ruta no implementa el flujo de abastecimiento.", "compras-simulada.png"),
            ("E10 - Configuracion simulada: la ruta no ofrece ajustes operativos reales.", "configuracion-simulada.png"),
        ],
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
