from __future__ import annotations

import sys
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\dev\SiCoSe")
SOURCE_PDF = Path(r"C:\Users\SAM\Downloads\ACT-5-Equipo5.pdf")
ASSET_DIR = ROOT / "output" / "word" / "consultorio_assets"
OUTPUT_DIR = ROOT / "output" / "word"
DOCX_PATH = OUTPUT_DIR / "AUDIT-1-Consultorio-Digital-SaaS-preliminar.docx"

SKILL_SCRIPTS = Path(
    r"C:\Users\SAM\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BASE_FONT = "Calibri"
ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(85, 85, 85)
SOFT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F8FAFC"
WHITE = "FFFFFF"

EVIDENCE_PAGES = [1, 2, 4, 5, 9, 10, 12, 14, 16]


def set_run_font(
    run,
    *,
    name: str = BASE_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell(cell) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10,
    color: RGBColor = INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def ensure_assets() -> list[Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if not SOURCE_PDF.exists():
        raise FileNotFoundError(SOURCE_PDF)

    document = pdfium.PdfDocument(str(SOURCE_PDF))
    generated = []
    for page_num in EVIDENCE_PAGES:
        out_path = ASSET_DIR / f"page-{page_num}.png"
        if not out_path.exists():
            page = document[page_num - 1]
            bitmap = page.render(scale=2)
            bitmap.to_pil().save(out_path)
        generated.append(out_path)
    return generated


def style_document(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = BASE_FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = ACCENT_DARK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_DARK, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    left = paragraph.add_run("AUDIT-1 | Consultorio Digital SaaS")
    set_run_font(left, size=9, color=RGBColor(99, 110, 125))
    paragraph.add_run("\t")
    right = paragraph.add_run("Pagina ")
    set_run_font(right, size=9, color=RGBColor(99, 110, 125))
    add_page_number_field(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size=9, color=RGBColor(99, 110, 125))


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    set_paragraph_format(kicker, before=0, after=2, line_spacing=1.0)
    run = kicker.add_run("AUDIT-1 | REVISION DOCUMENTAL PRELIMINAR")
    set_run_font(run, size=10, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    set_paragraph_format(title, before=0, after=4, line_spacing=1.0)
    run = title.add_run("Consultorio Digital SaaS")
    set_run_font(run, size=26, color=ACCENT_DARK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, before=0, after=10, line_spacing=1.05)
    run = subtitle.add_run(
        "Evaluacion documental del MVP publicado en Vercel y del material tecnico ACT-5-Equipo5"
    )
    set_run_font(run, size=13, color=MUTED)


def add_label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.7, 4.8], 9360)
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.rows[index].cells[0], label, bold=True, size=10, color=ACCENT_DARK)
        shade_cell(table.rows[index].cells[0], HEADER_FILL)
        set_cell_text(table.rows[index].cells[1], value, size=10, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_light_callout(doc: Document, headline: str, body: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=8, line_spacing=1.05)
    shade_paragraph(paragraph, LIGHT_FILL)
    paragraph.paragraph_format.left_indent = Inches(0.05)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    run = paragraph.add_run(f"{headline} ")
    set_run_font(run, size=10.75, color=INK, bold=True)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.75, color=INK)


def add_overview_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.6, 2.0, 2.9], 9360)
    headers = ["Elemento", "Descripcion", "Lectura QA"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)

    rows = [
        ("Producto", "SaaS para consultorio/clínica", "Se entiende con rapidez."),
        ("Usuarios", "Secretaria, doctor y paciente", "Roles bien definidos."),
        ("Canal publico", "Landing + login de acceso", "Entrada simple y clara."),
        ("Canal privado", "Dashboard por rol", "Arquitectura coherente."),
        ("Autenticacion", "JWT + refresh token + OTP", "Flujo serio y realista."),
        ("Backend", "Supabase + Edge Functions", "Reduce simulacion superficial."),
    ]
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[idx], text, size=9.5, color=INK, align=alignment)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_criteria_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.6, 1.2, 3.8], 9360)
    headers = ["Criterio", "Resultado preliminar", "Sustento"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)

    rows = [
        ("Se entiende", "Si", "La documentacion explica roles, flujos, seguridad e infraestructura."),
        ("Flujo principal documentado", "Si", "OTP, citas, consulta y refresh token estan descritos de extremo a extremo."),
        ("Evidencia de no simulacion", "Si", "Supabase, Edge Functions, RLS y ABAC apuntan a logica real de backend."),
        ("Validacion viva en este entorno", "No disponible", "El dominio no respondio desde el entorno de trabajo; no se inventaron resultados."),
    ]
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=9.5, color=INK, align=alignment)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_risk_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.4, 2.7, 1.0, 2.4], 9360)
    headers = ["Area", "Riesgo / prueba", "Sev.", "Verificacion requerida"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=9.5, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)

    rows = [
        ("OTP", "Confirmar proveedor real y expiracion del codigo", "Alta", "Prueba de envio, expiracion, reintentos y bloqueo por abuso."),
        ("Sesion", "Validar refresh token rotation y almacenamiento seguro", "Alta", "Inspeccion de cookies, 401 y renovacion automatica."),
        ("Acceso", "Probar RBAC, ABAC y RLS con cuentas distintas", "Alta", "Intentos cruzados entre doctor, secretaria y paciente."),
        ("API", "Confirmar que keys no viajan al frontend", "Media", "Revision del bundle y variables de entorno."),
        ("CI/CD", "Verificar que existan tests y despliegue reproducible", "Media", "Evidencia de pipeline, build y tests automatizados."),
    ]
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], text, size=9.25, color=INK, align=alignment)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)


def add_image_figure(doc: Document, image_path: Path, caption: str) -> None:
    figure = doc.add_paragraph()
    set_paragraph_format(figure, before=0, after=2, line_spacing=1.0)
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.add_run().add_picture(str(image_path), width=Inches(6.0))

    caption_paragraph = doc.add_paragraph()
    set_paragraph_format(caption_paragraph, before=0, after=10, line_spacing=1.0)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9.25, color=MUTED, italic=True)


def add_cover_page(doc: Document) -> None:
    add_title_block(doc)

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=10, line_spacing=1.08)
    run = intro.add_run(
        "Este reporte resume la lectura tecnica del proyecto y documenta una validacion preliminar. "
        "No se inventaron pruebas vivas: el entorno no permitio acceso al sitio publico, por lo que el dictamen se apoya en la documentacion tecnica disponible."
    )
    set_run_font(run, size=11, color=INK)

    add_label_value_table(
        doc,
        [
            ("Link evaluado", "https://consultorio-digital-saas.vercel.app/"),
            ("Repositorio", "https://github.com/SoyAri/consultorio-digital"),
            ("Documento base", str(SOURCE_PDF)),
            ("Estado", "Revision documental preliminar, sin validacion viva"),
            ("Fecha", "18 de junio de 2026"),
        ],
    )

    add_light_callout(
        doc,
        "Alcance:",
        "se revisaron el proposito del producto, su arquitectura, el plan de pruebas, la seguridad y los flujos de negocio descritos. "
        "El anexo visual contiene capturas del documento tecnico base para dejar trazabilidad del analisis.",
    )


def add_summary_page(doc: Document) -> None:
    heading = doc.add_heading("1. Resumen ejecutivo", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "Consultorio Digital se describe como un SaaS para consultas medicas con tres perfiles principales: secretaria/admin, doctor y paciente. "
        "La documentacion muestra un MVP con logica realista: landing publica, acceso diferenciado por rol, OTP para pacientes, dashboard privado, citas, historiales, seguridad por tokens y control de acceso por reglas de negocio."
    )
    set_run_font(run, size=11, color=INK)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "Desde la optica QA, el material no se ve como una maqueta superficial. Hay coherencia entre frontend, backend, base de datos, seguridad y plan de pruebas. "
        "El unico limite de esta entrega es metodologico: el dominio publico no fue accesible desde este entorno, asi que la validacion no puede marcarse como viva."
    )
    set_run_font(run, size=11, color=INK)

    add_overview_table(doc)


def add_methodology_page(doc: Document) -> None:
    heading = doc.add_heading("2. Metodologia y alcance", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    set_paragraph_format(intro, before=0, after=8, line_spacing=1.08)
    run = intro.add_run(
        "La revision se hizo como auditoria documental preliminar. Se analizaron las paginas 1 a 16 del PDF de arquitectura para reconstruir la propuesta de valor, los flujos de usuario y los controles tecnicos planeados."
    )
    set_run_font(run, size=11, color=INK)

    bullets = [
        "Visión general, público objetivo y módulos principales.",
        "Stack tecnologico y justificacion (Angular, Supabase, PostgreSQL, Vercel/Netlify, JWT y RBAC/ABAC).",
        "Modelo de datos, autenticacion y gestion de sesion.",
        "Seguridad, infraestructura, flujo del sistema y plan de pruebas.",
    ]
    for bullet in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(paragraph, before=0, after=4, line_spacing=1.08)
        run = paragraph.add_run(bullet)
        set_run_font(run, size=10.75, color=INK)

    add_criteria_table(doc)

    add_light_callout(
        doc,
        "Lectura QA:",
        "el producto tiene una narrativa clara y suficiente detalle tecnico para sustentar un MVP serio. "
        "La validacion de produccion o staging queda pendiente de acceso vivo y evidencia directa del sitio.",
    )


def add_technical_page(doc: Document) -> None:
    heading = doc.add_heading("3. Lectura tecnica del MVP", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run("Arquitectura observada en la documentacion:")
    set_run_font(run, size=11, color=INK, bold=True)

    architecture_points = [
        ("Presentacion", "Angular 21 como SPA con modulos lazy-loaded y rutas separadas por rol."),
        ("Lógica", "Supabase con Edge Functions para OTP, refresh token, citas y notificaciones."),
        ("Datos", "PostgreSQL con RLS activado y modelo de entidades para usuarios, pacientes, citas e historiales."),
        ("Acceso", "RBAC para roles y ABAC para restricciones por doctor/paciente."),
        ("Infraestructura", "Vercel/Netlify para hosting y GitHub Actions para CI/CD."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = column_widths_from_weights([1.7, 5.8], 9360)
    headers = ["Componente", "Lectura tecnica"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10, color=ACCENT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], HEADER_FILL)
    for component, detail in architecture_points:
        row = table.add_row()
        set_cell_text(row.cells[0], component, size=9.5, color=INK, bold=True)
        set_cell_text(row.cells[1], detail, size=9.5, color=INK)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=8, after=4, line_spacing=1.08)
    run = paragraph.add_run(
        "En QA esto se interpreta como una base con intención real de implementación, no como demo vacía. "
        "La parte que sigue pendiente no es el diseño, sino la verificacion de que el dominio publico efectivamente ejecute ese comportamiento."
    )
    set_run_font(run, size=11, color=INK)

    add_risk_table(doc)


def add_conclusion_page(doc: Document) -> None:
    heading = doc.add_heading("4. Conclusión y dictamen preliminar", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "La documentacion revisada describe un MVP consistente, bien segmentado y con un enfoque tecnico maduro. "
        "Se entiende el producto, el flujo principal esta claro y el disenio de seguridad apunta a un sistema real, no a una simple maqueta."
    )
    set_run_font(run, size=11, color=INK)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run(
        "Dictamen: revision preliminar positiva a nivel documental, con validacion viva pendiente por restriccion de acceso a red en este entorno."
    )
    set_run_font(run, size=11, color=INK, bold=True)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=0, after=8, line_spacing=1.08)
    run = paragraph.add_run("Recomendacion QA:")
    set_run_font(run, size=11, color=INK, bold=True)
    run = paragraph.add_run(
        "cuando tengas acceso al entorno, ejecutar pruebas vivas de login, OTP, citas, refresco de sesion, RLS/ABAC y navegacion entre roles con capturas reales."
    )
    set_run_font(run, size=11, color=INK)

    add_light_callout(
        doc,
        "Score preliminar:",
        "3/4 en claridad, logica y robustez documentada. El punto no validado es la ejecucion viva contra el dominio publico.",
    )


def add_evidence_page(doc: Document, image_path: Path, heading_text: str, caption: str) -> None:
    heading = doc.add_heading(heading_text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_image_figure(doc, image_path, caption)


def build_doc() -> Path:
    asset_paths = ensure_assets()
    asset_map = {path.name: path for path in asset_paths}

    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0])
    doc.settings.odd_and_even_pages_header_footer = False
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    add_cover_page(doc)
    doc.add_page_break()
    add_footer(doc.sections[0])

    add_summary_page(doc)
    doc.add_page_break()
    add_methodology_page(doc)
    doc.add_page_break()
    add_technical_page(doc)
    doc.add_page_break()
    add_conclusion_page(doc)

    evidence_specs = [
        ("page-1.png", "Anexo 1 - Portada y contexto", "A1 - Portada del documento tecnico base, con autoria del equipo y titulo del proyecto."),
        ("page-2.png", "Anexo 2 - Vision general", "A2 - Vision general, publico objetivo y arquitectura de alto nivel."),
        ("page-4.png", "Anexo 3 - Frontend y rutas", "A3 - Mapa de pantallas y control de acceso por modulo."),
        ("page-5.png", "Anexo 4 - Modelo de datos", "A4 - Base de datos, entidades y relaciones principales."),
        ("page-9.png", "Anexo 5 - Autenticacion y sesion", "A5 - Flujo de login, OTP y refresh token."),
        ("page-10.png", "Anexo 6 - Control de acceso", "A6 - RBAC, ABAC y politicas RLS por tabla."),
        ("page-12.png", "Anexo 7 - Seguridad e infraestructura", "A7 - Controles de seguridad, variables de entorno e infraestructura."),
        ("page-14.png", "Anexo 8 - Flujo del sistema", "A8 - Registro, citas y consulta medica descritos paso a paso."),
        ("page-16.png", "Anexo 9 - Plan de pruebas", "A9 - Plan de pruebas unitarias, integracion, E2E y seguridad."),
    ]

    for image_name, heading_text, caption in evidence_specs:
        doc.add_page_break()
        add_evidence_page(doc, asset_map[image_name], heading_text, caption)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
